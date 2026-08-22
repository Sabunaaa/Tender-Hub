"""Extract searchable text from ტექნიკური დავალება attachments."""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Iterable

log = logging.getLogger(__name__)

SPEC_MARKER = "ტექნიკური"
MAX_FILES = 4
MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_CHARS = 200_000
SUPPORTED_EXT = {".pdf", ".xlsx", ".xls", ".docx"}

# Structure markers understood by the frontend spec panel. Spreadsheet rows keep
# their cells so the panel can rebuild the table instead of showing a text wall.
# Control characters, because real spec lines do start with things like "# ".
FILE_PREFIX = "\x01"
SHEET_PREFIX = "\x02"
CELL_SEP = "\t"

# Any whitespace except tab and newline, so cell and line boundaries survive.
_INLINE_SPACE = re.compile(r"[^\S\t\n]+")


def is_spec_filename(name: str | None) -> bool:
    return SPEC_MARKER in (name or "")


def spec_files(attachments: Iterable[dict]) -> list[dict]:
    """Unique URL list for attachments whose filename contains ტექნიკური."""
    seen: set[str] = set()
    out: list[dict] = []
    for att in attachments:
        name = att.get("name") or ""
        url = att.get("url") or ""
        if not url or url in seen or not is_spec_filename(name):
            continue
        ext = Path(name).suffix.lower()
        if ext and ext not in SUPPORTED_EXT:
            continue
        seen.add(url)
        out.append({"name": name, "url": url})
        if len(out) >= MAX_FILES:
            break
    return out


def _cell(value: object) -> str:
    """One spreadsheet/table cell flattened to a single line."""
    text = str(value).replace("\t", " ").replace("\n", " ")
    return _INLINE_SPACE.sub(" ", text).strip()


def _row(cells: list[str]) -> str:
    """Tab-joined row, or "" when every cell is blank."""
    while cells and not cells[-1]:
        cells.pop()
    return CELL_SEP.join(cells) if cells else ""


def _paragraph(text: str) -> str:
    return _INLINE_SPACE.sub(" ", text.replace("\t", " ")).strip()


def _join(lines: Iterable[str]) -> str:
    """Drop blank lines and collapse the result into one text block."""
    return "\n".join(line for line in lines if line.strip())


def extract_bytes(data: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _from_pdf(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    if ext == ".xls":
        return _from_xls(data)
    if ext == ".docx":
        return _from_docx(data)
    return ""


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    lines: list[str] = []
    for page in reader.pages:
        # Layout mode keeps each table row on one line; plain mode splits a row
        # into one fragment per cell, which reads as noise. Layout mode gives up
        # on pages with rotated text and returns nothing, so fall back per page.
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except Exception:
            text = ""
        if not text.strip():
            text = page.extract_text() or ""
        for line in text.splitlines():
            lines.append(_paragraph(line))
    return _join(lines)


def _sheet_lines(name: str, rows: Iterable[list[str]], multi_sheet: bool) -> list[str]:
    body = [row for row in (_row(cells) for cells in rows) if row]
    if not body:
        return []
    header = [f"{SHEET_PREFIX}{_paragraph(name)}"] if multi_sheet else []
    return header + body


def _from_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    try:
        sheets = wb.worksheets
        for sheet in sheets:
            rows = (
                [_cell(c) if c is not None else "" for c in row]
                for row in sheet.iter_rows(values_only=True)
            )
            lines.extend(_sheet_lines(sheet.title, rows, len(sheets) > 1))
    finally:
        wb.close()
    return _join(lines)


def _from_xls(data: bytes) -> str:
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    sheets = book.sheets()
    for sheet in sheets:
        rows = (
            [_cell(c) if c not in (None, "") else "" for c in sheet.row_values(r)]
            for r in range(sheet.nrows)
        )
        lines.extend(_sheet_lines(sheet.name, rows, len(sheets) > 1))
    return _join(lines)


def _from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    lines = [_paragraph(p.text) for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(_row([_cell(c.text) for c in row.cells]))
    return _join(lines)


def extract_spec_text(client, attachments: Iterable[dict]) -> str:
    """Download matching files through the portal session and concatenate text.

    Returns ``""`` when nothing could be extracted so callers can mark the
    tender as attempted and skip it next time.
    """
    chunks: list[str] = []
    for att in spec_files(attachments):
        try:
            data = client.get_bytes(att["url"])
        except Exception as exc:
            log.warning("spec download failed (%s): %s", att["name"], exc)
            continue
        if not data or len(data) > MAX_FILE_BYTES:
            log.warning("spec skipped (%s): empty or too large (%s bytes)", att["name"], len(data or b""))
            continue
        try:
            text = extract_bytes(data, att["name"])
        except Exception as exc:
            log.warning("spec parse failed (%s): %s", att["name"], exc)
            continue
        if text:
            chunks.append(f"{FILE_PREFIX}{_paragraph(att['name'])}\n{text}")
    return "\n".join(chunks)[:MAX_CHARS]
